"""
Generates a 20-document fictional 'Contoso Corp' HR/benefits/procedures
corpus across four formats (md, html, docx, pdf) for Use Case 1.

Run once: python generate_corpus.py
Outputs into corpus/generated/{markdown,html,docx,pdf}/
"""
import textwrap
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).parent / "generated"
(ROOT / "markdown").mkdir(parents=True, exist_ok=True)
(ROOT / "html").mkdir(parents=True, exist_ok=True)
(ROOT / "docx").mkdir(parents=True, exist_ok=True)
(ROOT / "pdf").mkdir(parents=True, exist_ok=True)

# Each doc: (title, [(heading, paragraphs[])])
DOCS_MARKDOWN = [
    ("Employee Handbook Introduction", [
        ("Welcome to Contoso Corp", [
            "This handbook describes the policies, benefits, and procedures that apply to all Contoso Corp employees, effective January 1, 2026. It is maintained by the People Operations team and reviewed annually, or sooner if a material change in law or company practice requires it. Every policy in this handbook has a designated owner within People Operations, Legal, or Finance, listed in the internal Policy Register alongside the date of last review.",
            "It supersedes all prior versions, including any handbook or policy summary distributed before January 1, 2026. Where a specific team's policy differs from this handbook, the team policy governs only if it has been approved in writing by HR and filed with the central Policy Register; unapproved team-level policies (for example, an informal manager practice around flexible hours) do not override this handbook and should not be relied upon in a dispute.",
            "Nothing in this handbook constitutes an employment contract or a guarantee of continued employment. Contoso Corp is an at-will employer in every jurisdiction where at-will employment is legally permitted, meaning either the employee or Contoso Corp may end the employment relationship at any time, with or without cause or notice, subject to any legal exceptions (such as protections against discriminatory termination) described elsewhere in this handbook.",
            "This handbook is written for a U.S. audience. Employees working from a Contoso Corp office or subsidiary outside the United States should refer to their country-specific addendum, which takes precedence over this handbook wherever the two conflict, since local labor law often mandates terms this general handbook does not address.",
            "Printed or PDF copies of this handbook may not reflect the most current version; the version published on the internal HR portal is always the authoritative one, and employees relying on a saved or printed copy older than 90 days should confirm they are not missing a recent update before citing it in any HR conversation.",
        ]),
        ("Who This Applies To", [
            "This handbook applies to all full-time and part-time employees of Contoso Corp in the United States, including employees on approved leave. Contractors, consultants, and interns should refer to their engagement agreement or internship agreement for applicable terms, since many handbook sections (PTO, 401(k), health benefits) do not extend to non-employee workers, and misapplying employee benefits to a contractor can create legal misclassification risk for Contoso Corp.",
            "Employees working under a collective bargaining agreement should refer to that agreement first; where the agreement is silent, this handbook applies as a supplement, not a replacement. In the rare event a specific handbook provision directly conflicts with a signed collective bargaining agreement, the agreement controls for the employees it covers.",
            "Newly acquired employees (for example, staff joining Contoso Corp through an acquisition) become subject to this handbook on the date their offer letter specifies, not on the date the acquisition itself closed; People Operations communicates this transition date individually to each affected employee.",
        ]),
        ("How to Use This Handbook", [
            "Each section is self-contained and cross-references related policies where relevant (for example, the Paid Time Off Policy references the Parental Leave Policy for leave that extends beyond standard PTO). If you cannot find an answer here, contact HR at hr@contoso-corp.example or ask the Knowledge Assistant, which is grounded in this same document set and will tell you honestly when a question falls outside what these policies cover.",
            "Policy updates are announced company-wide by email at least 2 weeks before taking effect, except where a change is required immediately by law, in which case the effective date and a plain-language summary of the change are still communicated as soon as practically possible.",
            "A version history of every policy change since January 1, 2026 is kept by People Operations and is available on request; this handbook itself does not display prior versions inline to keep it readable, but nothing is changed silently.",
        ]),
        ("Handbook Governance and Amendments", [
            "Proposed changes to this handbook originate from People Operations, Legal, or Finance depending on the policy area, and go through a review committee before publication. Employees may submit a suggestion for a policy clarification or change through the internal Policy Feedback form; suggestions are reviewed quarterly, though not every suggestion results in a change.",
            "Emergency amendments (for example, a policy change required by a sudden change in state law) may bypass the standard quarterly review cycle, but still require sign-off from both Legal and the VP of People Operations before publication, and are always accompanied by a clear notice explaining why the normal notice period was shortened.",
            "This handbook is published in English as the authoritative version. A Spanish-language translation is maintained for convenience and reviewed for accuracy at each major update, but in the event of any discrepancy between the two versions, the English version governs.",
        ]),
        ("Where to Get Help", [
            "For a question about pay or benefits enrollment, contact HR directly rather than your manager, since managers do not have access to payroll or benefits systems. For a question about a specific policy's intent or history, People Operations can provide context the written policy text alone may not capture.",
            "The Knowledge Assistant referenced throughout this handbook is intended for quick lookups and is grounded strictly in this document set; for anything involving a personal or sensitive situation (a leave request tied to a health condition, a harassment concern), employees should go directly to HR or the Ethics Hotline rather than relying on the assistant.",
        ]),
        ("Frequently Asked Questions", [
            "Does this handbook apply to me if I was hired through a staffing agency? No — agency-placed workers are covered by their agency's own policies, not this handbook, even while working on-site at a Contoso Corp office; if a Contoso Corp manager directs an agency worker's day-to-day work, that manager should still confirm expectations with the agency directly rather than applying this handbook informally.",
            "What happens if a policy in this handbook conflicts with something my manager told me verbally? The written handbook, as published on the internal HR portal, always controls over any verbal statement from a manager, regardless of the manager's seniority; a verbal exception a manager grants is only valid if it is also confirmed in writing by HR.",
            "Who decides if a new situation isn't clearly covered by any policy here? People Operations makes the initial call, in consultation with Legal for anything with potential legal exposure; a genuinely novel situation not covered by any existing policy typically results in a documented one-off decision that later informs whether the handbook itself should be updated.",
        ]),
        ("Related Policies and Scope", [
            "This handbook does not cover compensation structure, equity grants, or sales commission plans, which are governed by separate, role-specific compensation documents distributed at hire and updated as needed; employees with compensation questions should contact their HR Business Partner rather than searching this handbook for an answer it does not contain.",
        ]),
        ("Key Definitions", [
            "'Full-time employee' means an employee regularly scheduled for 30 or more hours per week. 'Part-time employee' means an employee regularly scheduled for fewer than 30 hours per week but at least 15. Employees scheduled for fewer than 15 hours per week are classified as 'variable-hour' and are not eligible for most benefits described in this handbook.",
            "'Continuous employment' means employment without a break in service; a leave of absence approved under any policy in this handbook (parental leave, medical leave, military leave) does not break continuous employment for purposes of calculating tenure-based eligibility (for example, PTO accrual rate or 401(k) vesting).",
        ]),
        ("Roles and Responsibilities", [
            "People Operations owns the content and accuracy of this handbook. Managers are responsible for applying policies consistently within their team and escalating any situation they are unsure how to handle rather than guessing. Employees are responsible for reading and following the policies that apply to their employment status.",
        ]),
    ]),
    ("Paid Time Off Policy", [
        ("Accrual", [
            "Full-time employees accrue 1.5 days of paid time off (PTO) per month, for a total of 18 days per year. Accrual is calculated on the last day of each calendar month and posted to the employee's PTO balance within 3 business days. For example, an employee who starts on March 10th accrues a prorated 0.55 days for March (21 remaining days in the month divided by 31, multiplied by 1.5), then accrues the full 1.5 days starting in April.",
            "Part-time employees accrue PTO on a pro-rated basis according to scheduled hours: an employee scheduled for 20 hours per week (50% of full-time) accrues 0.75 days per month, and an employee scheduled for 30 hours per week (75% of full-time) accrues 1.125 days per month. Scheduled hours are reviewed quarterly, and an employee's accrual rate is adjusted going forward (not retroactively) if their scheduled hours change.",
            "PTO accrual begins on an employee's first day of employment; there is no waiting period. However, newly accrued PTO cannot be used until an employee has completed 30 days of employment, except in the case of a documented medical emergency, in which case HR may approve early use of accrued-but-unvested PTO on a case-by-case basis.",
            "Employees returning from an unpaid leave of absence resume accrual on their return-to-work date; PTO does not accrue during the unpaid portion of a leave, but does continue to accrue during any paid portion (for example, the paid weeks of parental leave described in the Parental Leave Policy).",
            "Employees who transfer internally between departments keep their existing PTO balance and accrual rate unchanged; a transfer is not treated as a new hire date for accrual purposes, since accrual is tied to total continuous Contoso Corp tenure, not tenure in any specific role or department.",
        ]),
        ("Rollover", [
            "Employees may roll over up to 5 unused PTO days into the following calendar year. Any balance above 5 days is forfeited on December 31st unless local law requires otherwise (for example, California and several other states prohibit use-it-or-lose-it PTO policies, and employees in those states carry over their full unused balance instead of being capped at 5 days).",
            "Rolled-over days must be used by March 31st of the new year or they are forfeited; the standard 5-day cap does not reset until that date, meaning an employee could theoretically hold up to 5 rolled-over days plus a growing new-year balance simultaneously between January 1st and March 31st.",
            "Employees who are on an approved leave of absence spanning the December 31st rollover date are granted an automatic 60-day extension on their rollover usage deadline (to May 31st instead of March 31st), to account for time they were unable to schedule PTO usage.",
        ]),
        ("Requesting Time Off", [
            "Submit PTO requests through the Contoso HR portal at least 5 business days in advance for planned absences. Same-day requests for illness do not require advance notice, but should be reported to your manager before your shift start time, or as soon as reasonably possible if the illness itself prevents timely notice.",
            "Requests during Contoso Corp's peak season (the last two weeks of December) require manager approval at least 3 weeks in advance and are granted on a first-requested, first-approved basis; teams with more than 40% of members requesting the same peak-season week may need to stagger approvals, which managers should communicate as early as possible.",
            "Managers must respond to a PTO request within 2 business days; a request that is neither approved nor denied within that window is automatically approved. If a manager denies a request, the denial must include a written reason (for example, insufficient team coverage) and the employee may escalate to their HR Business Partner if they believe the denial was applied inconsistently compared to teammates.",
        ]),
        ("Unused PTO at Termination", [
            "Upon voluntary or involuntary termination, Contoso Corp pays out accrued and unused PTO at the employee's final base rate of pay, in accordance with state law. This payout is included in the employee's final paycheck, calculated as the number of unused accrued days multiplied by the employee's daily base pay rate at the time of termination.",
            "Employees terminated for cause forfeit any rolled-over PTO from a prior year but still receive payout for PTO accrued in the current calendar year, where required by state law. Employees who resign voluntarily receive payout for all accrued and unused PTO, including any properly rolled-over balance, regardless of the reason for resignation.",
        ]),
        ("Sick Leave", [
            "Sick leave is drawn from the same PTO bank described above rather than a separate sick-leave allotment, except in states that legally require a distinct, separately tracked sick-leave balance (for example, several states mandate a minimum number of paid sick days independent of any general PTO policy). In those states, Contoso Corp maintains a separate statutory sick-leave balance alongside the standard PTO balance, and employees in those states can see both balances listed separately in the HR portal.",
            "Extended illness beyond what standard PTO can reasonably cover should be discussed with HR about short-term disability coverage or FMLA leave, both of which are described in separate policies; PTO is not intended to cover a multi-week medical absence on its own.",
        ]),
        ("Bereavement and Jury Duty Leave", [
            "Employees receive up to 5 paid bereavement days for the death of an immediate family member (spouse, child, parent, sibling) and up to 2 paid days for an extended family member, in addition to standard PTO; bereavement days do not draw down the employee's PTO balance.",
            "Employees summoned for jury duty receive paid leave for the full duration of their service, with no cap on the number of days, upon providing the court summons to HR; employees called as a witness in a legal proceeding unrelated to Contoso Corp business should discuss the specific situation with HR, since coverage depends on the nature of the proceeding.",
        ]),
        ("Frequently Asked Questions", [
            "Can I take PTO during my introductory 30-day period? You accrue PTO from day one, but cannot use it until 30 days have passed, except for a documented medical emergency approved by HR; plan any vacation requests for after your 30-day mark.",
            "What if I need more time off than I have accrued? Unpaid leave beyond your accrued PTO balance requires manager approval and, depending on the reason and duration, may need to be processed as a formal leave of absence rather than extended PTO; talk to your manager and HR before assuming unpaid time off will be approved informally.",
            "Does a company holiday that falls during my PTO count against my PTO balance? No — if a recognized Contoso Corp holiday falls within a scheduled PTO block, that day is paid as a holiday and is not deducted from the employee's PTO balance.",
        ]),
        ("Related Policies", [
            "This policy should be read alongside the Parental Leave Policy for any leave related to birth, adoption, or foster placement, and alongside the Remote Work Policy's core-hours section for how PTO interacts with a remote work schedule spanning multiple time zones.",
        ]),
        ("Key Definitions", [
            "'Accrued PTO' means the portion of an employee's annual PTO allotment they have earned to date based on months of service, as opposed to their full annual entitlement. 'Rolled-over PTO' means unused accrued PTO carried into the next calendar year under the Rollover section above, tracked separately from the new year's fresh accrual.",
        ]),
        ("Roles and Responsibilities", [
            "Employees are responsible for submitting requests with adequate notice and tracking their own balance in the HR portal. Managers are responsible for responding to requests within the required window and applying approval criteria consistently across their team. Payroll is responsible for accurate accrual calculation and final payout at termination.",
        ]),
    ]),
    ("Remote Work Policy", [
        ("Eligibility", [
            "Employees whose role has been designated 'remote-eligible' by their department head may work remotely up to 5 days per week. Roles requiring on-site equipment, in-person client contact, or physical presence for security/compliance reasons (e.g., handling regulated hardware) are not remote-eligible, regardless of the employee's preference or tenure.",
            "Remote eligibility is reviewed at each performance cycle and can be revoked if an employee's role responsibilities change to require more on-site presence; a revocation must be communicated at least 30 days before it takes effect, except where a client contract or regulatory requirement forces an immediate change.",
            "Hybrid arrangements (for example, 2 days remote and 3 days in-office) are set at the department level, not the individual level, so that a manager can plan team collaboration days consistently; individual exceptions require the department head's written approval.",
        ]),
        ("Equipment", [
            "Contoso Corp provides a laptop, monitor, and a one-time $300 home office stipend for remote-eligible employees, disbursed through payroll within the first full pay cycle after a remote work agreement is signed. Requests for additional equipment (a second monitor, ergonomic chair, keyboard) go through the IT Service Desk and are evaluated case by case, typically approved within 5 business days for requests under $150.",
            "Employees are responsible for maintaining a secure and adequately private workspace; company equipment used at home remains company property and must be returned upon separation, per the Offboarding Procedure. Equipment lost or damaged due to employee negligence (as opposed to normal wear or a documented theft with a police report) may be charged back to the employee at replacement cost.",
            "Internet reimbursement of $50 per month is available to remote-eligible employees working 3 or more remote days per week; employees working fewer remote days are not eligible for this specific stipend but remain eligible for the one-time equipment stipend.",
        ]),
        ("Core Hours", [
            "Remote employees must be reachable during core hours of 10:00 AM to 3:00 PM in their local time zone, Monday through Friday, excluding company holidays. Outside of core hours, employees may structure their schedule flexibly as long as they meet their weekly hour commitment and attend any meetings their manager designates as mandatory.",
            "Employees working across more than one time zone from their team (e.g., a distributed team spanning Eastern and Pacific time) should agree on a shared core-hours overlap with their manager in writing, typically a minimum 3-hour daily overlap window covering both time zones' business hours.",
            "'Reachable' means responsive to chat, email, or a scheduled call within a reasonable window (generally under 2 hours) during core hours — it does not require being continuously active online, and managers may not use presence-indicator status (e.g., an 'away' chat icon) alone as a basis for a performance concern.",
        ]),
        ("Working From Outside the Country", [
            "Working from outside the United States for more than 20 consecutive business days requires prior approval from HR and Legal due to tax and employment law implications. Requests must be submitted at least 30 days before the intended travel date using the International Remote Work Request Form.",
            "Approval is not guaranteed and depends on whether Contoso Corp has a registered business presence or applicable tax treaty in the destination country; some countries are excluded entirely due to sanctions or data residency restrictions, and the current excluded-country list is maintained by Legal and available on the internal Compliance portal.",
            "Working from outside the country for 20 consecutive business days or fewer does not require this approval process, but employees should still notify their manager for coverage-planning purposes and confirm with IT that company equipment will function reliably given local network conditions.",
        ]),
        ("Home Workspace Safety and Security", [
            "Remote employees are expected to maintain a workspace free of obvious safety hazards (exposed wiring, blocked exits) and are covered by workers' compensation for injuries that occur during work hours and are directly related to work tasks, in the same way an office-based injury would be covered. A brief self-certification of workspace safety is required annually through the HR portal.",
            "Confidential data accessed while remote must follow the same rules as the Data Security and Confidentiality Policy — for example, a screen displaying customer data should not be visible through a window or to household members, and company video calls involving confidential topics should use a virtual background or a private room where reasonably available.",
        ]),
        ("Coworking Spaces and Alternate Locations", [
            "Employees may work from a coworking space or coffee shop occasionally, but should avoid discussing confidential business matters audibly in a public space, and should never leave a company laptop unattended and unlocked. Coworking space membership fees are not reimbursed under this policy; employees who want a dedicated coworking membership should discuss it with their manager as a separate budget request.",
        ]),
        ("Frequently Asked Questions", [
            "Can I permanently relocate to a different state while remaining remote-eligible? Yes, but you must notify HR and Payroll at least 30 days before the move, since a state change can affect tax withholding, benefits eligibility, and in some cases whether your role remains remote-eligible if the new state has licensing or regulatory implications for your specific role.",
            "Is remote eligibility the same as a formal flexible-work accommodation for a medical condition? No — a medical accommodation goes through the separate Accommodation Requests process described in the Anti-Harassment and Non-Discrimination Policy and can grant remote work even for a role not otherwise designated remote-eligible, since it is evaluated under different legal standards.",
            "What happens to my equipment stipend if I switch from remote-eligible to hybrid partway through the year? The one-time $300 equipment stipend is not repeated or clawed back; the $50 monthly internet stipend stops as of the pay period following the schedule change, prorated for the transition month.",
        ]),
        ("Roles and Responsibilities", [
            "Department heads are responsible for designating which roles are remote-eligible and reviewing that designation at each performance cycle. IT is responsible for provisioning and supporting remote equipment. Employees are responsible for maintaining a secure, reasonably professional workspace and meeting the core-hours reachability standard.",
        ]),
    ]),
    ("Code of Conduct", [
        ("Professional Behavior", [
            "Contoso Corp employees are expected to treat colleagues, clients, and vendors with respect and professionalism at all times, both in person and in digital communication, including internal chat tools, email, and video calls. This expectation applies equally at company social events and off-site gatherings that are organized or sponsored by Contoso Corp.",
            "Disagreements about work product or approach are expected and healthy; personal attacks, public humiliation, or retaliatory behavior in response to disagreement are not tolerated and are grounds for disciplinary action up to termination. Managers receive specific training on how to distinguish constructive critical feedback from conduct that crosses this line.",
        ]),
        ("Conflicts of Interest", [
            "Employees must disclose any financial or personal relationship that could reasonably be seen as influencing a business decision, using the Conflict of Interest Disclosure Form, filed with Compliance within 10 business days of the relationship arising or being recognized. Failure to disclose a known conflict, even if the underlying relationship itself was permissible, is treated as a separate and more serious violation than the conflict itself.",
            "Common examples requiring disclosure: a close family member employed by a vendor Contoso Corp is evaluating, a personal financial stake in a competitor, or a side business that could be seen as competing with Contoso Corp's services. Compliance reviews each disclosure and, in most cases, resolves it with a documented mitigation plan (for example, recusing the employee from vendor selection decisions) rather than requiring the employee to end the relationship.",
        ]),
        ("Gifts and Entertainment", [
            "Employees may accept gifts from vendors or clients valued under $75. Gifts above this threshold must be declined or reported to your manager and Compliance within 5 business days of receipt; reported gifts over $75 are typically redirected to a company-wide raffle or charitable donation rather than kept personally.",
            "Cash or cash-equivalent gifts (gift cards, cryptocurrency) of any amount must always be declined, with no exception for value, since cash-equivalents are considered inherently higher-risk regardless of the stated dollar amount.",
        ]),
        ("Reporting Concerns", [
            "Concerns about conduct violations can be reported confidentially through the Contoso Ethics Hotline at 1-800-555-0142 or ethics@contoso-corp.example. Reports may be made anonymously, and the hotline is operated by an independent third party, not by Contoso Corp's own HR team, so that employees can report concerns involving HR itself without a conflict.",
            "Retaliation against a good-faith reporter is strictly prohibited and is itself grounds for termination, regardless of the reporter's role or seniority relative to the person retaliating; this protection applies even when the underlying report is ultimately not substantiated, as long as it was made honestly and in good faith.",
        ]),
        ("Social Media Use", [
            "Employees are free to discuss their work at Contoso Corp on personal social media, but may not share confidential or non-public information (unreleased product details, financial results before public announcement, customer names) under any circumstances. Personal social media posts should include a disclaimer that views expressed are the employee's own when discussing anything work-related, particularly for employees in customer-facing or public-facing roles.",
            "Employees representing Contoso Corp on official company social media accounts must follow the separate Brand and Communications guidelines maintained by Marketing, which include a stricter approval process for public statements than this general policy requires for personal accounts.",
        ]),
        ("Political and Outside Activity", [
            "Employees may engage in political activity and personal outside employment on their own time, as long as it does not create a conflict of interest (see the Conflicts of Interest section) and does not use Contoso Corp resources, branding, or work time. Contoso Corp does not endorse any candidate or party as a company, and employees should not imply company endorsement when engaging in personal political activity.",
        ]),
        ("Frequently Asked Questions", [
            "Can I be disciplined for something I post on my personal social media outside of work hours? Only if the post violates this Code of Conduct directly — for example, harassing a coworker, disclosing confidential information, or making a discriminatory statement identifiable as coming from a Contoso Corp employee. Purely personal opinions unrelated to work or colleagues are not subject to this policy.",
            "What if I'm unsure whether a gift crosses the $75 threshold? When in doubt, disclose it to your manager and Compliance; disclosing a gift that turns out to be under the threshold carries no consequence, while failing to disclose one that turns out to be over the threshold is treated as a more serious violation than the gift itself.",
            "Is a report to the Ethics Hotline really anonymous if I want it to be? Yes — the hotline is operated by an independent third party specifically so that fully anonymous reports are possible; Contoso Corp only receives the substance of the report, not identifying information, unless the reporter chooses to provide it.",
        ]),
        ("Enforcement", [
            "Violations of this Code of Conduct are addressed through Contoso Corp's standard progressive discipline process for most infractions (verbal warning, written warning, final warning, termination), except for severe violations (violence, harassment, theft, deliberate data breach) which may result in immediate termination on a first offense, bypassing the earlier steps entirely.",
        ]),
    ]),
    ("Parental Leave Policy", [
        ("Eligibility", [
            "All full-time employees are eligible for paid parental leave after 90 days of continuous employment, regardless of gender or the method of family growth (birth, adoption, or surrogacy). Part-time employees working at least 20 hours per week are eligible for a pro-rated version of the same benefit, calculated at the same percentage used for their PTO accrual rate.",
            "Foster placements that are reasonably expected to last 6 months or longer also qualify under this policy, treated the same as an adoption placement for purposes of leave duration and pay.",
        ]),
        ("Duration and Pay", [
            "Birthing parents receive 16 weeks of fully paid leave, which can begin up to 2 weeks before the expected due date if medically necessary. Non-birthing parents and adoptive parents receive 10 weeks of fully paid leave.",
            "Leave must be taken within 12 months of the birth or placement date and may be split into two blocks with manager approval — for example, 6 weeks immediately after birth and the remaining 4 weeks taken later in the first year, which many parents use to cover a gap before daycare enrollment begins.",
            "Parental leave runs concurrently with any leave required under the Family and Medical Leave Act (FMLA) where the employee is FMLA-eligible; it does not stack on top of FMLA leave to extend total job-protected time beyond what FMLA and this policy separately provide. Employees who are not FMLA-eligible (for example, due to insufficient tenure) still receive the full paid leave duration under this Contoso Corp policy, just without the separate federal job-protection guarantee FMLA provides.",
        ]),
        ("Benefits Continuation", [
            "Health, dental, and vision benefits continue uninterrupted during parental leave at the employee's normal contribution rate; Contoso Corp continues to pay its employer share as if the employee were actively working, for the full duration of the paid leave.",
            "401(k) contributions pause during unpaid portions of leave (if any) but resume automatically upon return to active payroll status, with no re-enrollment required; since parental leave under this policy is fully paid for its entire duration, 401(k) contributions in practice do not pause at all unless an employee takes additional unpaid leave beyond what this policy covers.",
        ]),
        ("Return-to-Work Transition", [
            "Employees returning from parental leave may request a gradual return schedule (for example, part-time hours for the first 2 weeks back) with manager approval; this is a courtesy accommodation, not a guaranteed right, and compensation during a gradual return is prorated to actual hours worked once the fully-paid leave period has ended.",
            "Contoso Corp provides a private lactation room in every office location, available on a first-come basis, and remote employees returning from leave are not required to use video during calls if doing so would require compromising the privacy needed for pumping breaks.",
        ]),
        ("Leave for Pregnancy Loss or Complications", [
            "Employees experiencing a pregnancy loss, stillbirth, or a serious pregnancy-related medical complication are eligible for up to 2 weeks of paid leave in addition to any leave otherwise available under this policy or short-term disability coverage, without needing to have reached the standard 90-day eligibility threshold described above.",
        ]),
        ("Frequently Asked Questions", [
            "Can I use accrued PTO to extend my parental leave beyond the paid weeks this policy provides? Yes, PTO can be used consecutively after the paid parental leave period ends, subject to normal PTO request and approval rules; many employees use this to extend an effective leave by an additional 1-3 weeks depending on their PTO balance.",
            "Do both parents get leave if we both work at Contoso Corp and are having a child together? Yes — each employee's eligibility and leave duration is determined individually based on their own role (birthing or non-birthing parent), not reduced or shared because both parents work at the same company.",
            "What if my adoption falls through after I've already started my leave? Leave taken in good-faith anticipation of a placement that does not finalize is not clawed back; HR will work with the employee individually on returning to work, since this is treated as a sensitive personal situation rather than a policy violation.",
        ]),
        ("Roles and Responsibilities", [
            "HR is responsible for confirming eligibility and processing leave paperwork. Payroll is responsible for ensuring leave is paid correctly and benefits contributions continue as described. Managers are responsible for planning coverage during an employee's leave without pressuring the employee about work while they are out.",
        ]),
    ]),
]

DOCS_HTML = [
    ("Expense Reimbursement Policy", [
        ("Submitting Expenses", [
            "Submit expense reports through the Contoso Expense portal within 30 days of the expense date. Reports submitted after 60 days will not be reimbursed except in extenuating circumstances approved in writing by Finance leadership, such as an employee's own medical emergency preventing timely submission.",
            "Each report requires a business justification, the associated project or client code, and an itemized receipt for any line item over the receipt threshold described below. Reports missing a project or client code are routed back to the submitter automatically and are not queued for approval until corrected.",
        ]),
        ("Receipts", [
            "Original itemized receipts are required for any single expense over $25. Credit card statements alone are not sufficient documentation, since they do not itemize what was purchased, which Finance needs in order to confirm the expense falls within policy (for example, distinguishing a reimbursable meal from a non-reimbursable bar tab on the same statement line).",
            "Digital receipts (photographed or forwarded email receipts) are accepted as long as the total amount, vendor name, and date are legible; blurry or partial photos are the single most common reason an expense report is sent back for resubmission.",
        ]),
        ("Meal Limits", [
            "Meal reimbursement is capped at $20 for breakfast, $25 for lunch, and $45 for dinner per person, including tax and tip, while traveling on approved company business. Alcohol is not reimbursable as part of a solo meal expense, and any alcohol line item on a solo meal receipt is deducted from the reimbursed total even if the overall receipt is under the meal cap.",
            "Meals with clients or prospects may be submitted under Business Entertainment rather than the standard meal caps, but require listing all attendees and the business purpose, and require manager pre-approval if the total exceeds $200 per person. Business Entertainment meals may include modest alcohol, unlike solo meals, as long as it is proportionate to the meal and not the primary expense.",
        ]),
        ("Non-Reimbursable Items", [
            "Contoso Corp does not reimburse alcohol for solo meals, traffic or parking fines, personal entertainment, or spa services, regardless of business context. This list is intentionally non-exhaustive; Finance reserves discretion to decline any expense that does not have a clear, direct business purpose even if it is not explicitly named here.",
            "Personal upgrades (seat upgrades not required for medical reasons, hotel room upgrades beyond standard) are also non-reimbursable and must be paid for out of pocket, with only the standard-rate portion reimbursed if a receipt shows a blended charge covering both the standard rate and the upgrade.",
        ]),
        ("Corporate Card Usage", [
            "Employees who travel more than twice per quarter are issued a Contoso Corp corporate card; the card should be used for all reimbursable business expenses when available, since corporate card transactions pre-populate the expense report and reduce the chance of a missing-receipt delay.",
            "Personal charges on a corporate card are prohibited even if the employee intends to reimburse Contoso Corp directly; repeated personal use of a corporate card, even when repaid, is grounds for card revocation and may be treated as a Code of Conduct violation depending on frequency and amount.",
        ]),
        ("International Expense Handling", [
            "Expenses incurred in a foreign currency are converted to U.S. dollars using the Concur portal's daily exchange rate as of the transaction date, not the date the expense report is submitted; employees should keep the original foreign-currency receipt even though the report displays the converted amount.",
        ]),
        ("Frequently Asked Questions", [
            "What if I lose a receipt for an expense over $25? Submit a Missing Receipt Affidavit through the Expense portal explaining the amount, vendor, and business purpose; this is accepted occasionally, but a pattern of repeated missing receipts from the same employee triggers a review by Finance.",
            "Can I expense a gift for a colleague's birthday or a team celebration? Small team morale expenses (a birthday cake, a team lunch) up to $15 per person are reimbursable under the Team Building sub-category; this is separate from the Business Entertainment category used for client meals.",
            "How long does reimbursement actually take once a report is approved? Approved reports are paid out on the next scheduled reimbursement cycle, which runs weekly; most employees see reimbursement within 5-7 business days of manager approval.",
        ]),
        ("Roles and Responsibilities", [
            "Employees are responsible for submitting accurate, timely, well-documented reports. Managers are responsible for reviewing reports for reasonableness and policy compliance before approving, not simply rubber-stamping every submission. Finance is responsible for final compliance review, payout processing, and identifying any pattern of policy misuse.",
        ]),
    ]),
    ("Travel Policy", [
        ("Booking Travel", [
            "All flights and hotels for business travel must be booked through the Contoso Travel Portal (powered by Concur) to ensure duty-of-care coverage, meaning Contoso Corp always knows where traveling employees are in the event of an emergency such as a natural disaster or civil disruption at the destination.",
            "Booking outside the portal is permitted only when the portal genuinely cannot fulfill the itinerary (e.g., a small regional carrier not listed), and requires notifying the Travel Desk within 24 hours of booking so the trip is still logged for duty-of-care purposes; unreported out-of-portal bookings may be denied reimbursement even if the travel itself was legitimate business travel.",
        ]),
        ("Airfare Class", [
            "Economy class is standard for flights under 6 hours. Premium economy is permitted for flights over 6 hours. Business class requires VP-level pre-approval and is only available for flights over 8 hours or for employees with a documented medical accommodation on file with HR.",
            "Connecting itineraries are evaluated on total travel time, not per-segment flight time, when determining which airfare class applies — a trip made up of two 4-hour segments with a layover, for a total of 9+ hours of travel, qualifies for the premium economy threshold even though no single segment exceeds 6 hours.",
        ]),
        ("Hotel Limits", [
            "Standard nightly hotel rate limits are $220 in major metro areas (New York, San Francisco, Chicago) and $160 elsewhere in the continental United States. Rates above these limits require manager approval in advance, noting the reason (e.g., conference-mandated hotel, no lower-cost option within a reasonable distance).",
            "The metro-area list is reviewed annually and may be expanded; the current authoritative list is always the one published on the Travel Desk's internal page, not any list printed in a slide deck or prior year's version of this policy.",
        ]),
        ("Ground Transportation", [
            "Rideshare and taxi are preferred over rental cars for trips under 3 days. Rental cars require a stated business justification in the expense report and should be booked at the economy or compact tier unless the trip involves transporting equipment or more than 2 colleagues.",
            "Rideshare or taxi expenses between an employee's home and the airport for business travel are reimbursable; the same trip for a personal vacation appended to business travel (a 'bleisure' trip) is only reimbursable for the portion of the itinerary tied to the actual business dates.",
        ]),
        ("Passports, Visas, and Travel Documents", [
            "Employees traveling internationally for business are responsible for maintaining a valid passport with at least 6 months of remaining validity beyond the trip's return date, as many countries deny entry otherwise. Visa application fees for business travel are reimbursable; Contoso Corp's Travel Desk can advise on typical visa processing times but does not guarantee approval, which remains at the discretion of the destination country.",
        ]),
        ("Travel Insurance and Emergency Assistance", [
            "All business travel booked through the Contoso Travel Portal automatically includes emergency medical and evacuation insurance at no cost to the employee; the insurance provider's 24/7 emergency line is printed on the travel itinerary confirmation and should be contacted first in a medical emergency abroad, before local emergency services if reasonably possible, since the insurer can coordinate care and cost directly.",
        ]),
        ("Frequently Asked Questions", [
            "Can I extend a business trip for personal time (a 'bleisure' trip)? Yes, as long as the additional personal days do not increase the cost of the airfare compared to a direct round trip on the business dates alone; any cost difference caused by the extended dates is the employee's responsibility, and this should be flagged to the Travel Desk when booking.",
            "What if my flight is delayed or cancelled and I need to book a new one on the spot? Contact the Travel Desk's after-hours emergency line first if reasonably possible, since they can often rebook within policy faster than a traveler booking independently; if an independent booking is unavoidable, keep all documentation and it will be reviewed for reimbursement under the same policy limits.",
            "Are conference registration fees covered under this Travel Policy? No — conference and training registration fees are reimbursed under the separate Learning and Development budget, not the Travel Policy, though travel and lodging to attend that conference does fall under this policy.",
        ]),
        ("Key Definitions", [
            "'Duty-of-care coverage' means Contoso Corp's ability, through the Travel Portal, to know a traveling employee's whereabouts and reach them in an emergency. A 'bleisure trip' means personal travel time appended to a business trip, reimbursable only for the business-dates portion of the itinerary.",
        ]),
    ]),
    ("Health Insurance Benefits Guide", [
        ("Plan Options", [
            "Contoso Corp offers three medical plan tiers: the PPO Plus plan, the PPO Standard plan, and the High-Deductible Health Plan (HDHP) with a company-funded Health Savings Account. All three plans use the same nationwide provider network, so switching plan tiers during open enrollment does not require switching doctors.",
            "The PPO Plus plan has the lowest deductible and highest monthly premium; the HDHP has the highest deductible but the lowest premium, offset partly by Contoso Corp's annual HSA contribution of $750 for employee-only coverage or $1,500 for family coverage, deposited in two installments in January and July.",
        ]),
        ("Company Contribution", [
            "Contoso Corp covers 80% of the employee-only premium and 60% of the premium for dependents, across all three plan tiers. The remaining premium is deducted pre-tax from each paycheck in equal installments across the plan year, meaning a mid-year plan change (following a qualifying life event) recalculates the remaining installments rather than the annual total.",
        ]),
        ("Enrollment Windows", [
            "New hires have 30 days from their start date to enroll; coverage begins on the first day of the month following enrollment. Open enrollment for all employees runs each year from November 1 through November 15, with changes effective January 1.",
            "An employee who misses the 30-day new-hire window without a qualifying life event must wait until the next open enrollment period and goes without Contoso Corp medical coverage in the interim, though they remain eligible to enroll in dental and vision on the standard timeline if those elections were not missed.",
        ]),
        ("Qualifying Life Events", [
            "Marriage, divorce, birth or adoption of a child, and loss of other coverage all qualify for a special enrollment period outside the standard windows, within 30 days of the event. Supporting documentation (marriage certificate, birth certificate, loss-of-coverage letter) must be submitted with the enrollment change.",
        ]),
        ("Prescription Drug Coverage", [
            "All three medical plans use the same four-tier prescription drug formulary: generic, preferred brand, non-preferred brand, and specialty. Generic drugs have the lowest copay across all three plan tiers, and employees are encouraged to ask their doctor about a generic alternative before a non-preferred brand is prescribed, since the cost difference to the employee can be substantial.",
            "Specialty medications (typically for chronic or complex conditions) require prior authorization and are filled through a designated specialty pharmacy rather than a standard retail pharmacy, which also provides additional clinical support for managing those medications.",
        ]),
        ("Telehealth Services", [
            "All three medical plans include unlimited telehealth visits for general medical concerns at a lower copay than an in-person visit, available 24/7 through the plan's telehealth app. Telehealth is not a substitute for emergency care, and the plan's telehealth provider will direct a caller to emergency services if the described symptoms warrant it.",
        ]),
        ("Frequently Asked Questions", [
            "Can I add my domestic partner to my medical plan even though we're not legally married? Yes — Contoso Corp recognizes domestic partnerships for benefits purposes upon completion of a Domestic Partner Affidavit, and the enrollment window and documentation requirements are the same as for a spouse.",
            "What happens to my HSA balance in the HDHP if I switch to a different plan tier next year? The HSA balance is the employee's own money and is never forfeited or clawed back regardless of plan changes; however, Contoso Corp's annual HSA contribution is only made while the employee is actively enrolled in the HDHP specifically.",
            "Is out-of-network care ever covered under the PPO plans? Yes, at a reduced reimbursement rate compared to in-network care, and subject to a separate, higher out-of-network deductible; the HDHP has more limited out-of-network coverage than the two PPO tiers.",
        ]),
        ("Roles and Responsibilities", [
            "HR administers plan enrollment and qualifying life event changes. The plan's insurance carrier administers claims and provider network questions directly; Contoso Corp does not process individual medical claims itself. Employees are responsible for reviewing plan documents and asking HR before assuming a specific service is covered.",
        ]),
    ]),
    ("Dental and Vision Benefits", [
        ("Dental Coverage", [
            "The Contoso dental plan covers preventive care (cleanings, exams, X-rays) at 100%, basic procedures (fillings) at 80%, and major procedures (crowns, root canals) at 50%, after a $50 annual deductible per covered individual. Preventive care does not count toward the deductible and is not limited even if a covered individual has already used other dental benefits earlier in the plan year.",
            "The plan has an annual maximum benefit of $2,000 per covered individual; costs above that maximum in a plan year are the employee's responsibility. Unused maximum benefit does not roll over to the following year.",
        ]),
        ("Vision Coverage", [
            "Vision coverage includes one eye exam per year at 100%, and a $150 allowance toward frames or contact lenses every 12 months. Designer frame upgrades beyond the allowance are available at a negotiated discount through the plan's partner retailers, typically 20-30% off retail price on the amount above the allowance.",
        ]),
        ("Orthodontia", [
            "Orthodontic treatment for dependents under 19 is covered at 50%, up to a lifetime maximum of $2,000 per covered individual. Adult orthodontia is not covered under the standard plan; employees seeking adult orthodontic coverage should ask HR about the optional voluntary dental buy-up plan, which is offered separately at full employee cost.",
        ]),
        ("Waiting Periods", [
            "New enrollees have no waiting period for preventive dental or vision care, but major dental procedures (crowns, root canals) have a 6-month waiting period from the enrollment date, waived for employees who can show 12 continuous months of prior dental coverage under a previous employer's plan.",
        ]),
        ("Out-of-Network Care", [
            "Both the dental and vision plans allow out-of-network care, but reimbursement is based on a set 'usual and customary' fee schedule rather than the provider's actual billed amount, meaning an employee choosing an out-of-network provider may owe a larger out-of-pocket balance than the same procedure would cost in-network.",
        ]),
        ("Frequently Asked Questions", [
            "Are teeth whitening or other cosmetic dental procedures covered? No, cosmetic-only procedures with no functional or health necessity are excluded from the dental plan entirely; a procedure with both a cosmetic and functional component (for example, a crown that is both structurally necessary and improves appearance) is covered based on the functional necessity, not the cosmetic benefit.",
            "Can I use my vision allowance on both glasses and contacts in the same 12-month period? No, the $150 allowance applies once per 12-month period toward either frames/lenses or contacts, not both; employees who want both should plan which to prioritize within a given benefit year.",
        ]),
        ("Key Definitions", [
            "'Preventive care' means routine cleanings, exams, and X-rays not tied to treating a specific problem. 'Major procedure' means a dental procedure like a crown or root canal that addresses significant structural damage or decay, as opposed to a 'basic procedure' like a filling.",
        ]),
    ]),
    ("401(k) Retirement Plan Summary", [
        ("Company Match", [
            "Contoso Corp matches 100% of employee 401(k) contributions up to the first 4% of eligible pay, and 50% of the next 2%, for a maximum company match of 5% when an employee contributes at least 6% of eligible pay. An employee contributing only 3% of pay receives a 3% match (the full 100% tier); an employee contributing 6% or more always receives the full 5% match regardless of how much above 6% they contribute.",
            "Matching contributions are calculated and deposited each pay period, not as a single annual true-up, so employees who front-load contributions early in the year should confirm they are not missing match dollars later in the year — Contoso Corp does perform a true-up calculation in January of the following year to correct any shortfall caused by front-loading and hitting the IRS annual contribution limit before year-end.",
        ]),
        ("Vesting Schedule", [
            "Employee contributions are always 100% vested. Company matching contributions vest over 3 years: 0% in year one, 50% in year two, and 100% after three years of service. Unvested matching funds are forfeited if an employee leaves before full vesting, and forfeited amounts are used to offset future plan administrative costs rather than redistributed to other participants.",
        ]),
        ("Eligibility and Enrollment", [
            "Employees become eligible to enroll in the 401(k) plan on the first day of the month following 30 days of employment, and are auto-enrolled at a 3% contribution rate unless they opt out or select a different rate within their first 45 days. Employees who take no action are auto-enrolled into the plan's default target-date fund, based on an assumed retirement age of 65.",
        ]),
        ("Investment Options", [
            "The plan offers a lineup of 18 index and actively managed funds spanning U.S. equity, international equity, bonds, and a suite of target-date funds. Employees can change their fund allocation at any time through the plan provider's website, with changes typically taking effect within 1-2 business days.",
        ]),
        ("Loans and Hardship Withdrawals", [
            "Active employees may borrow up to 50% of their vested balance (maximum $50,000) as a 401(k) loan, repaid through payroll deduction over up to 5 years at a fixed interest rate set at loan origination; an unpaid loan balance becomes a taxable distribution if the employee separates from Contoso Corp before it is repaid.",
            "Hardship withdrawals are permitted for a narrow set of IRS-defined reasons (preventing eviction, certain medical expenses, funeral expenses) and require documentation; unlike a loan, a hardship withdrawal is not repaid and is subject to income tax and, in most cases, an early-withdrawal penalty if the employee is under 59 and a half.",
        ]),
        ("Frequently Asked Questions", [
            "If I leave Contoso Corp, what happens to my 401(k) balance? Your own contributions and any vested match remain yours; you can leave the balance in the plan (if above a minimum threshold), roll it into an IRA or a new employer's plan, or take a cash distribution subject to taxes and an early-withdrawal penalty if you are under 59 and a half.",
            "Can I contribute more than the standard IRS annual limit if I'm over 50? Yes, employees aged 50 and over may make additional catch-up contributions above the standard IRS limit, and the plan automatically permits this once the system detects an employee's date of birth qualifies.",
            "Does Contoso Corp's match apply to catch-up contributions? No, the company match formula applies only up to the percentages described above on eligible pay; catch-up contributions do not receive an additional match beyond that.",
        ]),
        ("Roles and Responsibilities", [
            "HR and Payroll administer enrollment, contribution elections, and auto-enrollment defaults. The plan's third-party recordkeeper administers individual account access, investment elections, loans, and withdrawals directly; Contoso Corp itself does not have access to make changes to an individual employee's account.",
        ]),
    ]),
]

DOCS_DOCX = [
    ("Employee Referral Program", [
        ("Referral Bonus Amounts", [
            "Employees who refer a successful candidate for a standard role receive a $1,500 referral bonus. Referrals for hard-to-fill engineering or leadership roles earn a $3,000 bonus, as designated on the current Hard-to-Fill Roles list maintained by Talent Acquisition and updated quarterly based on open-requisition aging data.",
        ]),
        ("Payout Timing", [
            "Referral bonuses are paid in two installments: 50% after the referred employee's 90th day, and the remaining 50% after their 6-month anniversary, provided both the referring employee and the referred employee remain active Contoso Corp employees at each payout date.",
            "If the referring employee leaves the company before a payout date, that installment is forfeited; it is not paid out to the departed employee retroactively, and it is not redirected to anyone else either — the referral bonus simply is not paid for that installment.",
        ]),
        ("Eligibility", [
            "All active full-time employees are eligible to refer candidates, except employees on the Talent Acquisition team (whose role already involves recruiting) and hiring managers referring into their own open requisitions, to avoid a conflict of interest. Part-time employees are eligible to refer but the bonus amount is not pro-rated; it is paid at the full standard or hard-to-fill amount regardless of the referring employee's own schedule.",
        ]),
        ("Rehire and Boomerang Referrals", [
            "Referring a former Contoso Corp employee (a 'boomerang' hire) is eligible for the same bonus structure as any other referral, provided the former employee left in good standing and at least 6 months have passed since their departure; referring a former employee who was terminated for cause is not eligible for a bonus under any circumstance.",
        ]),
        ("Program Review and Fraud Prevention", [
            "Talent Acquisition audits referral bonus payouts quarterly to confirm the referring employee had a genuine prior relationship with the candidate; a referral submitted with no evidence of a real connection (for example, submitted the same day a resume was already in the applicant tracking system from a job board) may be denied pending review.",
        ]),
        ("Frequently Asked Questions", [
            "Can I refer more than one candidate for the same role? Yes, there is no limit on the number of referrals a single employee can submit, and each successful referral that results in a hire earns its own separate bonus, subject to the same eligibility and payout rules.",
            "What counts as a 'successful' referral for bonus purposes? A referral is successful once the referred candidate is hired and completes their first day; the bonus payout schedule (50% at 90 days, 50% at 6 months) then begins from that first day, regardless of how long the hiring process itself took.",
        ]),
        ("Roles and Responsibilities", [
            "Talent Acquisition maintains the Hard-to-Fill Roles list and audits payouts for legitimacy. Payroll processes the actual bonus disbursement on the schedule above. Referring employees are responsible for submitting the referral before the candidate applies independently, since only referrals submitted first are credited.",
        ]),
    ]),
    ("Performance Review Process", [
        ("Review Cadence", [
            "Contoso Corp conducts formal performance reviews twice per year: a mid-year check-in in June that is developmental and not tied to compensation, and an annual review in December that informs compensation decisions for the following year.",
        ]),
        ("Rating Scale", [
            "Employees are rated on a 4-point scale: Exceeds Expectations, Meets Expectations, Partially Meets Expectations, and Does Not Meet Expectations. Ratings are based on both what was achieved (goals) and how it was achieved (values and collaboration), and a manager must provide at least one specific example supporting any rating other than Meets Expectations.",
        ]),
        ("Calibration", [
            "Manager ratings go through a cross-team calibration session before being finalized, to ensure consistent application of the rating scale across departments and prevent rating inflation or deflation in any single team.",
            "Employees may request a rating explanation meeting with their manager within 5 business days of receiving a final rating, particularly if the calibrated rating differs from the manager's initial proposed rating; HR Business Partners attend these meetings on request if the employee prefers a neutral third party present.",
        ]),
        ("Performance Improvement Plans", [
            "An employee rated Does Not Meet Expectations, or Partially Meets Expectations for two consecutive review cycles, is placed on a formal Performance Improvement Plan (PIP) lasting 30 to 90 days, with specific, measurable goals agreed upon by the employee and manager and reviewed jointly with HR.",
            "Successful completion of a PIP removes the employee from formal monitoring but does not erase the underlying review history; failure to meet PIP goals typically results in termination, though HR reviews each case individually before that outcome is finalized.",
        ]),
        ("Promotion Nominations", [
            "Promotions are nominated by a manager during the annual review cycle and go through the same cross-team calibration process as ratings, to ensure consistent promotion standards across departments; an employee may also self-nominate for promotion consideration, which their manager is required to bring to calibration even if the manager does not personally support it.",
        ]),
        ("Frequently Asked Questions", [
            "If I'm rated Meets Expectations every cycle, can I still be promoted? Yes — a consistent Meets Expectations rating combined with demonstrated readiness for the next level's scope of work is a legitimate basis for promotion; a higher rating is not a strict prerequisite, though it does strengthen a promotion case during calibration.",
            "Can my rating be changed after calibration if I disagree with it? The calibration outcome is final for that review cycle, but a documented, substantive disagreement can be raised through the rating explanation meeting and, if unresolved, escalated to the HR Business Partner; this can inform how the *next* cycle's calibration is run, even though it does not retroactively change a already-finalized rating.",
        ]),
        ("Key Definitions", [
            "'Calibration' means the cross-team session where manager-proposed ratings are compared and adjusted for consistency before being finalized. A 'rating explanation meeting' is a follow-up conversation an employee may request to understand the reasoning behind their final calibrated rating.",
        ]),
    ]),
    ("Anti-Harassment and Non-Discrimination Policy", [
        ("Policy Statement", [
            "Contoso Corp prohibits discrimination or harassment based on race, color, religion, sex, national origin, age, disability, sexual orientation, gender identity, or any other status protected by applicable federal, state, or local law. This applies to hiring, promotion, compensation, and every other term of employment, and extends to interactions with clients, vendors, and job applicants, not only among employees.",
        ]),
        ("Reporting Process", [
            "Employees who experience or witness harassment should report it to their manager, HR Business Partner, or the confidential Ethics Hotline immediately. Reports are investigated within 10 business days by a trained investigator who is not in the reporting employee's direct chain of command, and both the reporting employee and the accused employee receive a summary of the outcome once the investigation concludes.",
        ]),
        ("Non-Retaliation", [
            "Contoso Corp strictly prohibits retaliation against anyone who reports a concern in good faith or participates in an investigation, even if the investigation does not substantiate the original complaint. Retaliation can include obvious actions like termination or demotion, but also subtler actions like exclusion from meetings or projects, and both are treated with equal seriousness under this policy.",
        ]),
        ("Third-Party Harassment", [
            "This policy also protects employees from harassment by a client, vendor, or other third party, not only by fellow Contoso Corp employees. An employee who experiences harassment from a client should report it the same way as internal harassment; Contoso Corp will address the situation with the client company directly, up to and including ending the business relationship if necessary.",
        ]),
        ("Accommodation Requests", [
            "Employees may request a reasonable accommodation for a disability, a sincerely held religious belief, or a pregnancy-related condition through HR's Accommodations Request Form; requests are evaluated on a case-by-case, interactive basis with the employee, and a request is denied only if it would impose an undue hardship on Contoso Corp's operations.",
        ]),
        ("Frequently Asked Questions", [
            "Can I report a concern about my own manager's behavior? Yes — the Ethics Hotline exists specifically so that concerns involving a direct manager, or anyone else in an employee's normal reporting chain, can be raised to someone outside that chain; reporting through the hotline does not require going through the manager in question first.",
            "What happens to the accused employee during a harassment investigation? Depending on the severity of the allegation, the accused employee may be placed on administrative leave (paid, not disciplinary) pending the outcome of the investigation, to prevent any possibility of continued harm or retaliation while the facts are being established.",
        ]),
        ("Enforcement", [
            "A substantiated violation of this policy results in discipline proportionate to severity, ranging from a documented warning to immediate termination for severe conduct; Contoso Corp does not apply a rigid one-size-fits-all penalty, since the appropriate response depends on the facts of each specific investigation.",
        ]),
    ]),
    ("IT Acceptable Use Policy", [
        ("Company Devices", [
            "Company-issued laptops and phones are for business use, with incidental personal use permitted as long as it does not interfere with work or violate other policies (for example, incidental personal use must not involve illegal content or excessive personal streaming that degrades network performance for others).",
        ]),
        ("Prohibited Activity", [
            "Employees may not install unauthorized software, disable endpoint security tools (antivirus, disk encryption, mobile device management), or use company systems to access illegal content. Violations may result in immediate loss of system access pending investigation, and repeated or willful violations are grounds for termination even on a first offense if the violation is severe enough (for example, deliberately disabling encryption to exfiltrate confidential data).",
        ]),
        ("Password Requirements", [
            "All Contoso accounts require multi-factor authentication and a password of at least 14 characters, rotated every 180 days. Password reuse across the last 10 passwords is blocked by the identity system, and passwords found in a known public breach database are automatically rejected at creation time regardless of how they otherwise score on length and complexity.",
        ]),
        ("Bring Your Own Device (BYOD)", [
            "Employees may access company email and calendar from a personal phone through the approved mobile management app, which enforces a device passcode and allows Contoso Corp to remotely wipe only the company data partition, not the entire personal device, if the phone is lost or the employee departs.",
            "Personal devices may not be used to access systems containing Confidential data (as defined in the Data Security and Confidentiality Policy); that level of access is restricted to company-issued and company-managed devices only.",
        ]),
        ("Generative AI Tool Usage", [
            "Employees may use approved generative AI tools for drafting and research, but must not paste Confidential or Internal data (customer information, unreleased financials, source code) into any AI tool that is not on the IT-approved list, since many public AI tools retain submitted input for model training by default.",
        ]),
        ("Frequently Asked Questions", [
            "Can I use my personal password manager for company account passwords? Yes, as long as the password manager itself supports multi-factor authentication; storing company passwords in an unsecured note, spreadsheet, or a password manager without MFA is a policy violation regardless of how strong the individual passwords themselves are.",
            "What should I do if I accidentally click a phishing link? Report it to security@contoso-corp.example immediately, even if you're not sure anything happened — do not wait to see if there are consequences first. Immediate reporting lets Security check for and contain any compromise quickly, and no employee is disciplined for reporting a mistake in good faith.",
        ]),
        ("Roles and Responsibilities", [
            "IT Security maintains endpoint protection tools and investigates policy violations. Employees are responsible for not circumventing security tools and for reporting suspected phishing or compromise promptly. Managers are responsible for ensuring their team completes required security training on schedule.",
        ]),
    ]),
    ("Data Security and Confidentiality Policy", [
        ("Data Classification", [
            "Contoso Corp classifies data into three tiers: Public, Internal, and Confidential. Customer PII and financial records are always classified Confidential, regardless of which system they are stored in, and this classification travels with the data even if it is copied or exported into a different tool.",
        ]),
        ("Handling Confidential Data", [
            "Confidential data may not be stored on personal devices, personal cloud storage, or transmitted over unencrypted channels. Confidential data leaving the company (e.g., shared with an approved third-party vendor) requires a signed data processing agreement on file with Legal before the first transfer occurs, not retroactively after data has already been shared.",
        ]),
        ("Incident Reporting", [
            "Suspected data breaches must be reported to security@contoso-corp.example within 1 hour of discovery, per the Incident Response Runbook. Delayed reporting can affect Contoso Corp's ability to meet legal breach-notification deadlines, so timeliness matters more than certainty — report suspected incidents even before they are confirmed, since the Security team, not the reporting employee, is responsible for determining whether a suspected incident is a genuine breach.",
        ]),
        ("Vendor Data Sharing", [
            "Sharing Confidential data with a third-party vendor requires a signed Data Processing Agreement reviewed by Legal, and the vendor must be added to Contoso Corp's approved vendor list before any data is transferred; this applies even to a well-known, reputable vendor, since the requirement is about the specific agreement in place, not the vendor's general reputation.",
        ]),
        ("Post-Employment Confidentiality", [
            "The confidentiality obligations in this policy continue after an employee leaves Contoso Corp, indefinitely for trade secrets and for a minimum of 3 years for other Confidential data, as outlined in the confidentiality agreement every employee signs at hire. Former employees remain bound by this even if their specific employment agreement is silent on the point, since this policy applies independently of the individual agreement.",
        ]),
        ("Frequently Asked Questions", [
            "Can I keep a copy of a presentation I made for a client after I leave Contoso Corp? Only the parts that do not contain Confidential client or company data; a generic template or your own individually-authored analysis approach may be retained if it does not disclose confidential specifics, but the underlying client data itself must not be retained.",
            "Does this policy apply to data about Contoso Corp employees themselves, not just customers? Yes — employee PII (social security numbers, compensation, medical information tied to a leave request) is classified Confidential under this policy exactly the same as customer data, and is subject to the same handling and incident-reporting rules.",
        ]),
        ("Key Definitions", [
            "'Confidential data' means data classified as such under the three-tier system described above (Public, Internal, Confidential). A 'data processing agreement' is a signed contract with a vendor governing how they may handle Confidential data shared with them, required before any such sharing occurs.",
        ]),
    ]),
]

DOCS_PDF = [
    ("Onboarding Checklist for New Hires", [
        ("Before Day One", [
            "IT provisions your laptop and accounts 3 business days before your start date. HR sends your offer packet and I-9 form for e-signature, which must be completed within 3 business days of the start date per federal requirements. New hires also receive a welcome email with parking or transit information for their assigned office, if applicable.",
        ]),
        ("Week One", [
            "New hires complete Contoso 101 orientation, benefits enrollment, and meet with their manager to set 30-60-90 day goals. IT security training and the Code of Conduct acknowledgment must also be completed within the first week, and access to systems beyond the basic starter set is granted only after both are confirmed complete.",
        ]),
        ("First 90 Days", [
            "Managers conduct check-ins at 30, 60, and 90 days. The first formal performance conversation happens at the 90-day mark, which also marks the end of the introductory period referenced in several other policies (e.g., 401(k) eligibility, parental leave eligibility, and full participation in the Employee Referral Program's eligibility rules).",
        ]),
        ("Buddy Program", [
            "Every new hire is assigned a peer buddy from a different team for their first 60 days, to provide an informal point of contact outside the direct reporting line for questions about culture and unwritten norms. Buddies are volunteers who have been at Contoso Corp for at least 1 year, and the buddy relationship is entirely optional to continue past the 60-day period if both people find it useful.",
            "Buddies receive a short onboarding guide of their own, covering common new-hire questions and where to redirect anything outside their own knowledge, so the buddy relationship supplements rather than replaces the manager and HR as the authoritative source for policy questions.",
        ]),
        ("Required Trainings", [
            "Beyond Contoso 101 and IT security training, new hires must complete anti-harassment training within their first 30 days and, for people managers specifically, a manager-fundamentals training within their first 90 days covering performance reviews, PTO approvals, and how to handle a conduct concern.",
            "Role-specific compliance training (for example, data-handling training for anyone touching Confidential data) is assigned automatically based on job code and must be completed within 2 weeks of assignment; overdue mandatory training triggers an automated reminder to both the employee and their manager.",
        ]),
        ("Manager Responsibilities", [
            "A new hire's manager is responsible for submitting the 30-60-90 day goals to HR, scheduling the check-ins described above, and confirming the new hire has been introduced to their assigned buddy within the first week. Managers who have not completed manager-fundamentals training may not yet approve PTO requests independently; HR co-approves in the interim.",
        ]),
        ("Frequently Asked Questions", [
            "What if my laptop hasn't arrived by my start date? Contact the IT Service Desk immediately; a loaner laptop is available for temporary use while your provisioned device is tracked down or replaced, so a shipping delay should never actually block your first day.",
            "Do I need to complete all required trainings before I can start real work? No — trainings have their own deadlines (30 days for anti-harassment, 2 weeks for role-specific compliance training) and are not a prerequisite for beginning normal job responsibilities, except for the specific systems access gated behind Week One's IT security training and Code of Conduct acknowledgment.",
        ]),
        ("Roles and Responsibilities", [
            "HR owns the offer packet, I-9 completion, and benefits enrollment. IT owns equipment provisioning and systems access. The hiring manager owns setting 30-60-90 day goals and conducting the scheduled check-ins. The assigned buddy owns informal cultural onboarding support.",
        ]),
        ("Office-Specific Orientation", [
            "New hires assigned to an office (as opposed to fully remote) receive a building tour on their first day covering emergency exits, restrooms, the kitchen and coffee area, printer and supply locations, and how badge access works for their specific floor. Employees should keep their badge on them at all times while in the building, since propping doors open or letting an unbadged person in behind them ('tailgating') is a security violation regardless of how well the new hire knows the person.",
            "Remote new hires receive an equivalent virtual orientation covering the same practical topics adapted for a home office context, plus a specific walkthrough of how to request the equipment and home office stipend described in the Remote Work Policy, since remote new hires do not automatically receive equipment the way an office-based new hire's desk is pre-set up.",
        ]),
        ("Feedback on the Onboarding Experience", [
            "At the 90-day mark, new hires receive a short onboarding experience survey, separate from their performance conversation, asking specifically about the checklist above: whether equipment arrived on time, whether the buddy program was useful, and whether required trainings were clear. This feedback is reviewed quarterly by People Operations and has directly driven changes to this checklist in the past, including the addition of the loaner-laptop process described in the FAQ above.",
        ]),
    ]),
    ("Offboarding Procedure", [
        ("Final Pay", [
            "Departing employees receive their final paycheck, including any unused PTO payout, on the next regularly scheduled payroll date after their last day, or sooner where required by state law (several states require final pay immediately or within 72 hours for involuntary terminations, and Payroll tracks each state's specific deadline in the offboarding checklist tool).",
        ]),
        ("Equipment Return", [
            "All company equipment (laptop, monitor, badge) must be returned within 5 business days of the last day of employment via prepaid shipping label provided by IT. Unreturned equipment may be deducted from final pay where permitted by state law, and IT remotely disables account access on the last day of employment regardless of whether equipment has yet been physically returned.",
        ]),
        ("Benefits After Departure", [
            "Health benefits end on the last day of the month of departure. COBRA continuation coverage information is mailed within 14 days, and the departing employee has 60 days from the COBRA notice to elect continuation coverage, which if elected is retroactive to the day after Contoso Corp coverage ended so there is no coverage gap.",
        ]),
        ("Exit Interview", [
            "Departing employees are invited to an optional exit interview with HR, conducted separately from the employee's manager, to gather candid feedback on their experience at Contoso Corp. Exit interview notes are anonymized and aggregated for department-level trend reporting; individual comments are not shared with the departing employee's former manager.",
            "Employees who decline the live exit interview may instead complete a short written survey covering the same topic areas; either format is entirely optional, and declining both has no effect on final pay, references, or eligibility for rehire.",
        ]),
        ("Rehire Eligibility", [
            "Employees who leave in good standing are generally eligible for rehire and may reapply for future openings at any time; their prior tenure does not automatically carry forward for benefits or PTO accrual purposes, since rehire is treated as a new hire date unless a specific rehire agreement states otherwise.",
        ]),
        ("Non-Disclosure Reminder", [
            "As part of offboarding, HR reviews the confidentiality obligations from the Data Security and Confidentiality Policy with every departing employee, which continue to apply after departure regardless of the reason for leaving. Departing employees are asked to confirm in writing that they have not retained any Confidential data on personal devices or accounts.",
        ]),
        ("Reference Requests", [
            "Contoso Corp provides only neutral employment verification (dates of employment and title) to external reference requests by default; a departing employee may request that their direct manager provide a more detailed personal reference, which is entirely the manager's individual choice to give, not a company-provided service.",
        ]),
        ("Frequently Asked Questions", [
            "Will I still get my final bonus or commission payout if I resign before the payout date? This depends on the specific compensation plan governing that bonus or commission, not this general Offboarding Procedure; check your compensation plan document or ask HR, since some plans require active employment on the payout date and others prorate based on time worked.",
            "Can I take my company laptop's personal files with me before returning it? IT can arrange a supervised transfer of clearly personal files (photos, personal documents) before wiping the device, but this must be scheduled in advance through the IT Service Desk and cannot happen after the device access has already been disabled.",
        ]),
        ("Roles and Responsibilities", [
            "HR coordinates the overall offboarding timeline and conducts the exit interview. IT disables access and processes equipment returns. Payroll calculates and issues final pay including PTO payout. The departing employee's manager is responsible for a knowledge-transfer plan so ongoing work is not disrupted.",
        ]),
        ("Knowledge Transfer Requirements", [
            "Departing employees, working with their manager, complete a Knowledge Transfer document covering active projects, key contacts, where important files and credentials live (without including actual passwords, which should never be written down), and any recurring task or deadline the employee normally owns. This is due before the employee's last day, and managers should build time for it into the employee's final two weeks rather than treating it as an afterthought.",
            "For roles with significant client-facing responsibility, the manager also schedules a formal handoff call with the affected client before the employee's last day, so the client hears about the transition directly rather than discovering it after the fact.",
        ]),
        ("Involuntary Termination Procedures", [
            "Involuntary terminations are conducted jointly by the manager and an HR Business Partner, never by a manager alone, to ensure consistency and legal compliance. The employee is provided a written notice of termination stating the effective date and, where applicable under state law, the reason. IT disables system access at the start of the termination meeting, not afterward, to reduce the risk of any data access issue during a sensitive conversation.",
        ]),
    ]),
    ("Workplace Safety Guidelines", [
        ("Office Safety", [
            "All Contoso offices maintain marked emergency exits, fire extinguishers inspected quarterly, and a designated floor warden for each office level responsible for coordinating evacuations during drills and real emergencies. Floor wardens are trained annually and their names are posted on each floor's safety board along with a backup warden in case the primary is out of office.",
        ]),
        ("Reporting Incidents", [
            "Any workplace injury, however minor, must be reported to Facilities and HR within 24 hours using the Incident Report Form, even if the employee does not seek medical treatment, so that any pattern of hazards can be identified early — for example, three separate minor slip reports near the same entrance in one month would trigger a facilities review of that specific area.",
        ]),
        ("Ergonomics", [
            "Employees may request a free ergonomic assessment of their home or office workstation through the Workplace Safety team. Assessments typically result in recommendations for chair height, monitor position, and keyboard placement, and may result in approval for an ergonomic equipment purchase beyond the standard remote work equipment stipend if the assessment identifies a genuine need.",
        ]),
        ("Emergency Procedures", [
            "In the event of a fire alarm, employees must evacuate immediately via the nearest marked exit and gather at their office's designated assembly point; re-entry is not permitted until the floor warden or emergency services confirm it is safe. Employees with a mobility limitation that affects evacuation should register with Facilities in advance so a specific evacuation plan (for example, using a designated evacuation chair) is prepared ahead of any actual emergency.",
            "Severe weather procedures (tornado, earthquake, or similar regional hazards) differ by office location and are posted on each floor's safety board separately from the standard fire evacuation route, since sheltering in place is often the correct response for these hazards rather than evacuating the building.",
        ]),
        ("Workplace Violence Prevention", [
            "Contoso Corp has zero tolerance for threats or acts of violence in the workplace, whether from an employee, client, or visitor. Any concern about a potential threat, even one that seems minor or ambiguous, should be reported immediately to Facilities Security and HR rather than waiting to see if the situation escalates.",
        ]),
        ("Hazardous Materials Handling", [
            "Offices with a print shop, lab space, or maintenance area maintain Safety Data Sheets for every chemical product on site, accessible through the Workplace Safety team's shared binder or portal. Employees who are not specifically trained and authorized should never attempt to clean up a chemical spill themselves; Facilities has a dedicated response protocol for this.",
        ]),
        ("Frequently Asked Questions", [
            "Do I need to report a near-miss (no injury, but something almost went wrong)? Yes — near-misses are specifically encouraged to be reported, since they often reveal a hazard before it actually causes an injury; the Incident Report Form has a specific near-miss option distinct from an actual injury report.",
            "Who pays for a workers' compensation claim from a workplace injury? Workers' compensation insurance, which Contoso Corp carries as required by law, covers medical costs and partial wage replacement for a qualifying workplace injury; the employee does not pay out of pocket for a covered claim, and using workers' compensation does not count against an employee's PTO balance.",
        ]),
        ("Roles and Responsibilities", [
            "Facilities maintains building safety equipment and coordinates floor wardens. The Workplace Safety team runs ergonomic assessments and safety training. Every employee is responsible for reporting hazards, injuries, and near-misses promptly, regardless of how minor they may seem.",
        ]),
        ("Visitor and Contractor Safety", [
            "Visitors and contractors must sign in at reception, wear a visible visitor badge at all times, and be escorted by their Contoso Corp host in any area beyond the main lobby and designated meeting rooms. Contractors performing maintenance or construction work are required to submit a safety plan to Facilities in advance for any work involving power tools, ladders above 6 feet, or temporary blocking of an emergency exit route.",
        ]),
        ("Security Cameras and Access Control", [
            "Common areas (lobbies, hallways, parking structures) are monitored by security cameras for safety purposes; footage is retained for 30 days and reviewed only in connection with a specific reported incident, not for general employee monitoring. Badge access logs are similarly retained and may be reviewed as part of a safety incident or security investigation.",
        ]),
    ]),
    ("Tuition Reimbursement Program", [
        ("Eligibility", [
            "Full-time employees with at least 6 months of tenure are eligible for tuition reimbursement for courses related to their current role or a reasonable career path at Contoso Corp, as determined jointly by the employee and their manager. Part-time employees are not eligible for this specific program, though they remain eligible for the separate professional development stipend described in the Learning and Development policy.",
        ]),
        ("Reimbursement Amount", [
            "Contoso Corp reimburses up to $5,250 per calendar year for tuition, required textbooks, and course fees, upon proof of a passing grade (C or better, or Pass in pass/fail courses). Amounts above $5,250 in a calendar year are reimbursed but treated as taxable income per IRS rules on employer education assistance, and Payroll will withhold applicable taxes from the portion above the threshold.",
        ]),
        ("Approval Process", [
            "Submit the Tuition Reimbursement Pre-Approval Form to your manager and HR before the course start date; reimbursement requests without prior approval will be denied regardless of the grade earned, since the pre-approval step is what confirms the course is genuinely job-related before Contoso Corp commits to reimbursing it.",
        ]),
        ("Repayment on Departure", [
            "Employees who voluntarily resign within 12 months of receiving tuition reimbursement must repay the reimbursed amount on a prorated basis, as outlined in the Tuition Reimbursement Agreement signed at approval time — for example, an employee who resigns 6 months after receiving reimbursement repays 50% of the amount received.",
            "Repayment is not required if the employee is involuntarily terminated without cause, or if they leave due to a qualifying medical hardship documented with HR; the repayment obligation exists specifically to discourage using the benefit and then immediately resigning, not to penalize departures outside the employee's control.",
        ]),
        ("Covered Program Types", [
            "Eligible programs include undergraduate and graduate degree programs, professional certifications directly relevant to the employee's role (for example, a cloud certification for an engineer), and accredited online courses; unaccredited bootcamps are reviewed case by case and are not automatically covered the way accredited coursework is.",
        ]),
        ("Denial Appeals", [
            "An employee whose pre-approval request is denied may appeal once, in writing, to a review panel made up of one HR representative and one representative from Learning and Development, who were not involved in the original decision; the panel's decision on appeal is final.",
        ]),
        ("Frequently Asked Questions", [
            "Can I use tuition reimbursement for a program outside the United States? Yes, as long as the institution is internationally accredited and the coursework meets the same current-role-or-reasonable-career-path standard applied to domestic programs; currency conversion for reimbursement uses the exchange rate on the date of your payment.",
            "What if I fail a single course in an otherwise multi-course program? Only the failed course is denied reimbursement; passing courses in the same program remain eligible, since eligibility is evaluated course by course, not for the program as a single unit.",
        ]),
        ("Roles and Responsibilities", [
            "The employee's manager confirms the course is relevant to the employee's role or reasonable career path. HR verifies eligibility and processes the pre-approval. Payroll disburses reimbursement upon proof of a passing grade and applies applicable tax withholding for amounts above the annual IRS-favorable threshold.",
        ]),
        ("Time Off for Coursework", [
            "Tuition reimbursement does not include paid time off to attend class or study; employees are expected to schedule coursework around work hours where possible, or use standard PTO for any classes that unavoidably conflict with the work schedule. Employees pursuing an employer-requested certification (as opposed to a personal career-development choice) may be granted paid study time at their manager's discretion, tracked separately from PTO.",
        ]),
        ("Program Funding Cap and Prioritization", [
            "The tuition reimbursement program has an overall annual company-wide budget; in the rare event that approved requests in a given year would exceed that budget, Learning and Development prioritizes requests tied to a documented business-critical skill gap over general career-development requests, though this cap has not been reached in any prior program year.",
        ]),
    ]),
    ("Employee Assistance Program Guide", [
        ("What the EAP Covers", [
            "The Contoso Employee Assistance Program (EAP) provides up to 6 free confidential counseling sessions per issue, per year, for employees and their household members. Sessions can be used for a range of concerns including stress, grief, relationship issues, and substance use, and a new 6-session allotment is available for each distinct issue, not a single shared pool across all concerns.",
        ]),
        ("Additional Services", [
            "The EAP also offers free legal consultations (up to 1 hour per matter), financial planning sessions, and childcare/eldercare referral services, all administered by the same third-party provider as the counseling benefit, so a single phone call can connect an employee to whichever of these services fits their situation.",
        ]),
        ("How to Access", [
            "Call the EAP support line at 1-800-555-0199, available 24/7, or visit the EAP portal linked from the Contoso benefits page. All contact is confidential and not shared with Contoso Corp, including whether a given employee has used the service at all — Contoso Corp only ever sees aggregate, anonymized utilization statistics, never individual usage records.",
        ]),
        ("Crisis Support", [
            "For an immediate mental health crisis, the EAP line offers a dedicated crisis option available at the same number, staffed by licensed clinicians who can also connect the caller to local emergency services if needed. This crisis option does not count against the standard 6-session annual allotment described above.",
            "Managers who become aware an employee may be in crisis are encouraged, but not required, to share the EAP crisis number directly; managers are never expected to assess or manage a mental health crisis themselves, and doing so is explicitly discouraged in favor of connecting the employee to the trained clinicians on the EAP line.",
        ]),
        ("Manager Referrals", [
            "A manager who is concerned about an employee's wellbeing based on observable workplace performance (not a diagnosis, which managers are not qualified to make) may suggest the employee consider the EAP; this suggestion must never be framed as a condition of continued employment, and using it that way is itself a policy violation.",
        ]),
        ("Household Member Eligibility", [
            "Household members eligible for EAP services include a spouse or domestic partner and any dependent living in the employee's home, including college-age dependents living away for school; the household member does not need to be enrolled in Contoso Corp's health insurance to be eligible for the EAP, since the two benefits are administered separately.",
        ]),
        ("Frequently Asked Questions", [
            "Will using the EAP show up anywhere in my personnel file? No — EAP usage is confidential and is never recorded in an employee's personnel file or shared with a manager; only aggregate, anonymized utilization statistics across the whole company are ever shared with Contoso Corp.",
            "Can I use the EAP for a work-related stress issue, or only personal matters? Work-related stress, conflict with a colleague, and burnout are all explicitly within scope for EAP counseling; the EAP is not limited to personal or family matters.",
        ]),
        ("Roles and Responsibilities", [
            "The third-party EAP provider delivers counseling and referral services directly and independently of Contoso Corp. HR communicates the benefit's existence and how to access it. Managers may suggest the EAP to a struggling employee but are never responsible for assessing or managing a mental health situation themselves.",
        ]),
        ("In-Person and Virtual Session Options", [
            "EAP counseling sessions are available both virtually (video or phone) and in person through the provider's local network of clinicians, depending on employee preference and local availability; virtual sessions are generally available faster than in-person appointments and are equally covered under the same 6-session allotment.",
        ]),
        ("Relationship to Short-Term Disability", [
            "The EAP is not a substitute for short-term disability coverage, which provides income replacement during an extended medical leave including for a mental health condition; an employee whose counseling needs extend beyond what the EAP's 6 free sessions can address should discuss short-term disability and continued treatment options with HR, who can explain how the two benefits work together.",
        ]),
    ]),
]


def write_markdown(title, sections):
    lines = [f"# {title}\n"]
    for heading, paras in sections:
        lines.append(f"## {heading}\n")
        for p in paras:
            lines.append(p + "\n")
    path = ROOT / "markdown" / f"{_slug(title)}.md"
    path.write_text("\n".join(lines), encoding="utf-8")


def write_html(title, sections):
    body = [f"<h1>{title}</h1>"]
    for heading, paras in sections:
        body.append(f"<h2>{heading}</h2>")
        for p in paras:
            body.append(f"<p>{p}</p>")
    html = f"<!doctype html><html><head><meta charset='utf-8'><title>{title}</title></head><body>{''.join(body)}</body></html>"
    path = ROOT / "html" / f"{_slug(title)}.html"
    path.write_text(html, encoding="utf-8")


def write_docx(title, sections):
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for heading, paras in sections:
        doc.add_heading(heading, level=2)
        for p in paras:
            run = doc.add_paragraph(p)
            run.style.font.size = Pt(11)
    path = ROOT / "docx" / f"{_slug(title)}.docx"
    doc.save(str(path))


def write_pdf(title, sections):
    path = ROOT / "pdf" / f"{_slug(title)}.pdf"
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    flow = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
    for heading, paras in sections:
        flow.append(Paragraph(heading, styles["Heading2"]))
        flow.append(Spacer(1, 6))
        for p in paras:
            flow.append(Paragraph(p, styles["BodyText"]))
            flow.append(Spacer(1, 6))
        flow.append(Spacer(1, 10))
    doc.build(flow)


def _slug(title: str) -> str:
    return title.lower().replace(" ", "_").replace("(", "").replace(")", "").replace("/", "-")


def _write_all(docs, writer_fn, label):
    """Write every document in a group independently — one locked/open
    file (e.g. a .docx open in Word) must never abort the rest of the run
    and leave the corpus in a half-old, half-new inconsistent state."""
    ok, failed = 0, []
    for title, sections in docs:
        try:
            writer_fn(title, sections)
            ok += 1
        except PermissionError:
            failed.append(title)
            print(f"  SKIPPED ({label}) '{title}' — file is open in another "
                  f"program (e.g. Word/Adobe). Close it and re-run this script.")
    return ok, failed


if __name__ == "__main__":
    results = [
        _write_all(DOCS_MARKDOWN, write_markdown, "markdown"),
        _write_all(DOCS_HTML, write_html, "html"),
        _write_all(DOCS_DOCX, write_docx, "docx"),
        _write_all(DOCS_PDF, write_pdf, "pdf"),
    ]

    total_ok = sum(ok for ok, _ in results)
    total_docs = len(DOCS_MARKDOWN) + len(DOCS_HTML) + len(DOCS_DOCX) + len(DOCS_PDF)
    all_failed = [title for _, failed in results for title in failed]

    print(f"Generated {total_ok}/{total_docs} documents: "
          f"{len(DOCS_MARKDOWN)} markdown, {len(DOCS_HTML)} html, "
          f"{len(DOCS_DOCX)} docx, {len(DOCS_PDF)} pdf")

    if all_failed:
        print(f"\n{len(all_failed)} document(s) were NOT updated (still old content) "
              f"because the file was open elsewhere:")
        for title in all_failed:
            print(f"  - {title}")
        print("Close those file(s) and re-run this script to finish updating them.")
