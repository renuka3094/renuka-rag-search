"""
Format-specific parsers. Each returns a list of (section_heading, text)
tuples so the chunker can stay structure-aware instead of treating every
file as one flat blob of text.
"""
from pathlib import Path

import markdown as md
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

Section = tuple[str | None, str]


def parse_pdf(path: Path) -> list[Section]:
    reader = PdfReader(str(path))
    sections: list[Section] = []
    for i, page in enumerate(reader.pages):
        text = (page.extract_text() or "").strip()
        if text:
            sections.append((f"Page {i + 1}", text))
    return sections


def parse_docx(path: Path) -> list[Section]:
    doc = DocxDocument(str(path))
    sections: list[Section] = []
    current_heading: str | None = None
    buffer: list[str] = []

    def flush():
        if buffer:
            sections.append((current_heading, "\n".join(buffer).strip()))
            buffer.clear()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name and para.style.name.lower().startswith("heading"):
            flush()
            current_heading = text
        else:
            buffer.append(text)
    flush()

    # Tables often carry policy details (e.g. PTO accrual tables) - include them.
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            sections.append((current_heading, "\n".join(rows)))

    return sections


def parse_html(path: Path) -> list[Section]:
    soup = BeautifulSoup(path.read_text(encoding="utf-8"), "html.parser")
    return _sections_from_soup(soup)


def parse_markdown(path: Path) -> list[Section]:
    html = md.markdown(path.read_text(encoding="utf-8"))
    soup = BeautifulSoup(html, "html.parser")
    return _sections_from_soup(soup)


def _sections_from_soup(soup: BeautifulSoup) -> list[Section]:
    sections: list[Section] = []
    current_heading: str | None = None
    buffer: list[str] = []

    def flush():
        if buffer:
            sections.append((current_heading, "\n".join(buffer).strip()))
            buffer.clear()

    for el in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "td"]):
        text = el.get_text(strip=True)
        if not text:
            continue
        if el.name in ("h1", "h2", "h3", "h4"):
            flush()
            current_heading = text
        else:
            buffer.append(text)
    flush()
    return sections


PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "html": parse_html,
    "htm": parse_html,
    "md": parse_markdown,
    "markdown": parse_markdown,
}


def parse_document(path: Path) -> list[Section]:
    ext = path.suffix.lower().lstrip(".")
    parser = PARSERS.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type: {ext}")
    return parser(path)
